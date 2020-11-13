import docx
import re

filename = 'template.docx'  # 模板文件,其中采用format方式书写模板
fileobj = docx.Document(filename)  # 模板文件对象
paranum = len(fileobj.paragraphs)  # 段落数量

content = ''  # word文件字符内容
content_paragraph = []

for para in fileobj.paragraphs:
    content += para.text+'\n'

pattern = '{.}'



